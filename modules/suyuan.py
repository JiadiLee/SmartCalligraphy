import gradio as gr
import os
import json
import sqlite3
import re
from pypinyin import lazy_pinyin, Style

try:
    from pypinyin import lazy_pinyin, Style
    PYPINYIN_AVAILABLE = True
except ImportError:
    PYPINYIN_AVAILABLE = False

from video import get_all_videos, search_videos, save_uploaded_video, get_video_by_id, update_video_views
from knowledge import search_knowledge, get_all_docs, add_knowledge_doc, rebuild_index, get_db_path, get_doc_videos
from knowledge_explain import get_personalized_learning_content
from config import get_storage_config


def get_docs_dir():
    storage_config = get_storage_config()
    base_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), storage_config.get('data_dir', 'data'))
    docs_dir = os.path.join(base_dir, 'documents')
    os.makedirs(docs_dir, exist_ok=True)
    return docs_dir


from pypinyin import lazy_pinyin, Style

from pypinyin import lazy_pinyin, Style

def sort_videos_by_title(videos):
    def get_sort_key(video):
        title = video[1]
        
        if not title:
            return (2, '') 

        first_char = title[0]

        if first_char.isascii() and first_char.isalnum():
            return (0, title.lower())

        try:
            context = title[:2]
            pinyin_list = lazy_pinyin(context, style=Style.FIRST_LETTER)
            if pinyin_list and pinyin_list[0].isalpha():
                return (1, pinyin_list[0].lower())
        except Exception as e:
            pass
            
        return (2, title)

    return sorted(videos, key=get_sort_key)

def sort_video_by_scores(video_scores):
    def get_sort_key(item):
        score = item[0]
        video_tuple = item[1]
        title = video_tuple[1]
        
        first_char = title[0] if title else ''
        pinyin_initial = 'z'
        
        if first_char.isascii() and first_char.isalnum():
            pinyin_initial = first_char.lower()
        else:
            try:
                context = title[:2]
                py_list = lazy_pinyin(context, style=Style.FIRST_LETTER)
                if py_list and py_list[0].isalpha():
                    pinyin_initial = py_list[0].lower()
            except:
                pass
        return (-score, pinyin_initial)

    return sorted(video_scores, key=get_sort_key)

VIDEO_CARDS_JS = """
<script>
function select_video(video_id, video_title) {
    var container = document.getElementById('video-cards-container');
    var dropdown = document.querySelector('.video-dropdown');
    var hiddenInput = document.querySelector('input[aria-label="选择要播放的视频"]');
    
    if (hiddenInput) {
        hiddenInput.value = video_id;
        hiddenInput.dispatchEvent(new Event('change'));
    }
    
    // Update dropdown
    var dropdownButton = document.querySelector('.video-dropdown button');
    if (dropdownButton) {
        dropdownButton.querySelector('span') && (dropdownButton.querySelector('span').textContent = video_title);
    }
}
</script>
"""


def refresh_video_choices():
    videos = get_all_videos(limit=50)
    videos = sort_videos_by_title(videos)
    choices = []
    for v in videos:
        v_id, v_title, _, _, _, _, _, v_views, _, _ = v
        choices.append((f"{v_title} (播放:{v_views})", str(v_id)))
    return gr.update(choices=choices)


def refresh_video_display():
    videos = get_all_videos(limit=50)
    videos = sort_videos_by_title(videos)
    html = build_video_cards(videos)
    choices = []
    default_value = None
    for i, v in enumerate(videos):
        v_id, v_title, _, _, _, _, _, v_views, _, _ = v
        choices.append((f"{v_title} (播放:{v_views})", str(v_id)))
        if i == 0:
            default_value = str(v_id)
    return html, gr.update(choices=choices, value=default_value)


def build_video_cards(videos, selected_id=None):
    if not videos:
        return "*暂无视频*"
    
    html = '<div id="video-cards-scroll-container" style="max-height: 400px; overflow-y: auto; padding: 5px;"><div style="display: grid; grid-template-columns: repeat(auto-fill, minmax(200px, 1fr)); gap: 15px; padding: 10px;">'
    
    for v in videos:
        v_id, v_title, v_desc, _, _, _, _, v_views, _, _ = v
        is_selected = 'border: 3px solid #4CAF50;' if str(v_id) == str(selected_id) else 'border: 1px solid #ddd;'
        
        html += f'''
        <div onclick="select_video('{v_id}', '{v_title} (播放:{v_views})')" 
             style="{is_selected} cursor: pointer; padding: 15px; border-radius: 8px; background: #f9f9f9; transition: all 0.2s;">
            <div style="font-weight: bold; font-size: 14px; margin-bottom: 8px; color: #333;">{v_title}</div>
            <div style="font-size: 12px; color: #666; overflow: hidden; text-overflow: ellipsis; display: -webkit-box; -webkit-line-clamp: 2; -webkit-box-orient: vertical;">
                {v_desc if v_desc else "无描述"}
            </div>
            <div style="font-size: 11px; color: #999; margin-top: 8px;">👁️ {v_views} 次播放</div>
        </div>
        '''
    
    html += '</div></div>'
    return html


def handle_select_video(video_id, video_title):
    return gr.update(value=video_id), gr.update(choices=[(video_title, video_id)], value=video_id)


def build_video_ui():
    all_videos = get_all_videos(limit=50)
    all_videos = sort_videos_by_title(all_videos)
    default_choices = []
    for v in all_videos:
        v_id, v_title, _, _, _, _, _, v_views, _, _ = v
        default_choices.append((f"{v_title} (播放:{v_views})", str(v_id)))
    
    video_cards_html = build_video_cards(all_videos)
    
    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### 📤 上传视频")
            
            upload_title = gr.Textbox(
                label="视频标题 *必填*",
                placeholder="请输入视频标题"
            )
            
            upload_desc = gr.Textbox(
                label="视频描述(可选)",
                placeholder="请输入视频描述",
                lines=2
            )
            
            upload_file = gr.File(
                label="选择视频文件",
                file_count="single"
            )
            
            upload_btn = gr.Button("📤 上传视频", variant="primary")
            upload_msg = gr.Markdown("")
        
        with gr.Column(scale=2):
            gr.Markdown("### 🔍 搜索视频")
            
            search_input = gr.Textbox(
                label="搜索关键词",
                placeholder="输入视频标题或描述搜索",
                lines=1
            )
            
            search_btn = gr.Button("🔍 搜索", variant="primary")
            
            video_list_display = gr.HTML(video_cards_html, elem_id="video-cards-container")
            
            selected_video_id = gr.Textbox(visible=False, value=default_choices[0][1] if default_choices else None, elem_id="selected-video-id")
    
    gr.Markdown("### ▶️ 播放视频")
    
    default_value = None
    if default_choices:
        default_value = default_choices[0][1]
    
    with gr.Row():
        video_dropdown = gr.Dropdown(
            label="选择要播放的视频",
            choices=default_choices,
            value=default_value,
            interactive=True,
            scale=3,
            elem_classes="video-dropdown"
        )
        with gr.Column(scale=1):
            play_btn = gr.Button("▶️ 播放", variant="primary", elem_classes="full-width-btn")
            delete_video_btn = gr.Button("🗑️ 删除", variant="stop", elem_classes="full-width-btn")
    
    with gr.Row():
        video_player = gr.Video(label="视频播放", visible=False, height=800, width="100%")
    
    def update_dropdown_from_js(video_id):
        if not video_id:
            return gr.update()
        all_v = get_all_videos(limit=50)
        for v in all_v:
            if str(v[0]) == str(video_id):
                v_title = v[1]
                v_views = v[7]
                return gr.update(choices=[(f"{v_title} (播放:{v_views})", str(v[0]))], value=str(v[0]))
        return gr.update()
    
    selected_video_id.change(
        fn=update_dropdown_from_js,
        inputs=[selected_video_id],
        outputs=[video_dropdown]
    )
    
    search_btn.click(
        fn=handle_video_search,
        inputs=[search_input],
        outputs=[video_list_display, video_dropdown, video_player]
    )
    
    upload_btn.click(
        fn=handle_video_upload,
        inputs=[upload_file, upload_title, upload_desc],
        outputs=[video_list_display, upload_msg, upload_title, upload_desc, upload_file, video_dropdown, video_player]
    )
    
    play_btn.click(
        fn=handle_play_video,
        inputs=[video_dropdown],
        outputs=[video_player]
    )
    
    delete_video_btn.click(
        fn=handle_delete_video,
        inputs=[video_dropdown],
        outputs=[video_list_display, video_player, video_dropdown, upload_msg]
    )
    
    return video_list_display, video_dropdown, upload_msg


def handle_video_search(keyword):
    if not keyword or not keyword.strip():
        all_videos = get_all_videos(limit=50)
        all_videos = sort_videos_by_title(all_videos)
        html = build_video_cards(all_videos)
        default_choices = []
        for v in all_videos:
            v_id, v_title, _, _, _, _, _, v_views, _, _ = v
            default_choices.append((f"{v_title} (播放:{v_views})", str(v_id)))
        default_value = default_choices[0][1] if default_choices else None
        return html, gr.update(choices=default_choices, value=default_value), gr.update(visible=False)
    
    videos = search_videos(keyword.strip())
    
    if not videos:
        return "*未找到匹配的视频*", gr.update(choices=[]), gr.update(visible=False)
    
    video_scores = []
    keyword_lower = keyword.strip().lower()

    for v in videos:
        v_id, v_title, v_desc, _, _, _, _, v_views, _, _ = v
        title_lower = v_title.lower() if v_title else ""
        desc_lower = v_desc.lower() if v_desc else ""
        
        score = 0
        if keyword_lower in title_lower:
            score += 100
            if title_lower.startswith(keyword_lower):
                score += 50
        if keyword_lower in desc_lower:
            score += 10
        
        video_scores.append((score, v))
    
    video_scores = sort_video_by_scores(video_scores)
    
    html = f'<div style="padding: 10px;"><p><b>📺 搜索结果 (共 {len(video_scores)} 个)</b></p></div>'
    html += build_video_cards([v for _, v in video_scores], str(video_scores[0][1][0]) if video_scores else None)
    
    choices = []
    default_value = None
    for i, (score, v) in enumerate(video_scores):
        v_id, v_title, v_desc, _, _, _, _, v_views, _, _ = v
        choice_label = f"{v_title} (播放:{v_views})"
        choices.append((choice_label, str(v_id)))
        
        if i == 0:
            default_value = str(v_id)
    
    return html, gr.update(choices=choices, value=default_value), gr.update(visible=False)
    

def handle_video_upload(video_file, title, description):
    if not title or not title.strip():
        return gr.update(), "⚠️ 请填写视频标题", gr.update(), gr.update(), gr.update(), gr.update(choices=[]), gr.update(visible=False)
    
    from video import check_video_title_exists
    if check_video_title_exists(title.strip()):
        return gr.update(), "⚠️ 视频名称已存在，请使用其他名称", gr.update(), gr.update(), gr.update(), gr.update(choices=[]), gr.update(visible=False)
    
    if not video_file:
        return gr.update(), "⚠️ 请选择视频文件", gr.update(), gr.update(), gr.update(), gr.update(choices=[]), gr.update(visible=False)
    
    try:
        save_uploaded_video(video_file, title.strip(), description.strip() if description else None)
        
        videos = get_all_videos(limit=50)
        videos = sort_videos_by_title(videos)
        
        html = build_video_cards(videos)
        
        choices = []
        default_value = None
        for i, v in enumerate(videos):
            v_id, v_title, v_desc, _, _, _, _, v_views, _, _ = v
            choice_label = f"{v_title} (播放:{v_views})"
            choices.append((choice_label, str(v_id)))
            if i == 0:
                default_value = str(v_id)
        
        return html, "✅ 视频上传成功！", gr.update(value=""), gr.update(value=""), gr.update(), gr.update(choices=choices, value=default_value), gr.update(visible=False)
    except Exception as e:
        return gr.update(), f"❌ 上传失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update(choices=[]), gr.update(visible=False)


def handle_play_video(video_id):
    if not video_id:
        return gr.update(visible=False)
    
    try:
        video = get_video_by_id(int(video_id))
        if video:
            update_video_views(int(video_id))
            return gr.update(visible=True, value=video[3])
    except:
        pass
    
    return gr.update(visible=False)


def handle_delete_video(video_id):
    if not video_id:
        return gr.update(), gr.update(visible=False), gr.update(choices=[]), "⚠️ 请先选择要删除的视频"
    
    try:
        from video import delete_video as db_delete_video
        success = db_delete_video(int(video_id))
        
        if success:
            videos = get_all_videos(limit=50)
            videos = sort_videos_by_title(videos)
            
            html = build_video_cards(videos)
            
            choices = []
            default_value = None
            for i, v in enumerate(videos):
                v_id, v_title, v_desc, _, _, _, _, v_views, _, _ = v
                choice_label = f"{v_title} (播放:{v_views})"
                choices.append((choice_label, str(v_id)))
                if i == 0:
                    default_value = str(v_id)
            
            return gr.update(value=html), gr.update(visible=False), gr.update(choices=choices, value=default_value), "✅ 视频已删除"
        else:
            return gr.update(), gr.update(visible=False), gr.update(choices=[]), "❌ 删除失败"
    except Exception as e:
        return gr.update(), gr.update(visible=False), gr.update(choices=[]), f"❌ 删除失败: {str(e)}"


def build_knowledge_ui():
    # 辅助函数：显示选中的文档详情
    def show_selected_doc(selected_value):
        if selected_value is None or selected_value == '':
            return format_all_documents()
        try:
            doc_id = selected_value
            if ':' in str(selected_value):
                doc_id = str(selected_value).split(':')[0]
            
            docs = get_all_docs()
            doc = None
            for d in docs:
                if d['id'] == doc_id:
                    doc = d
                    break
            
            if not doc:
                return format_all_documents()
            
            md = f"### 📄 文档详情\n\n"
            md += f"**ID: {doc['id']}**\n"
            md += f"**标题: {doc['title']}**\n"
            md += f"**难度: {doc['difficulty']}**\n"
            if doc.get('tags'):
                md += f"**标签: {', '.join(doc['tags'])}**\n"
            
            linked_videos = get_doc_videos(doc['id'])
            if linked_videos:
                md += f"\n**📺 关联视频：**\n"
                for v in linked_videos:
                    md += f"- {v['title']}\n"
            
            md += f"\n**内容：**\n{doc['content']}\n"
            return md
        except Exception as e:
            return f"❌ 加载失败: {str(e)}"
    
    # 先创建所有组件
    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### 📁 上传文档")
            with gr.Accordion("📄 上传文件", open=True):
                doc_file = gr.File(label="选择文件 (支持 txt, md, docx, pdf)", file_count="single", file_types=[".txt", ".md", ".docx", ".pdf"])
                gr.Markdown("---", visible=False)
                manual_title = gr.Textbox(label="文档标题 *必填*", placeholder="请输入文档标题", visible=False)
                manual_content = gr.Textbox(label="文档内容 *必填*", placeholder="请输入或粘贴文档内容", lines=6, visible=False)
                with gr.Row():
                    gr.Markdown("**关联视频 (可选):**")
                video_select = gr.Dropdown(label="选择关联视频", choices=[], interactive=True)
                add_doc_btn = gr.Button("➕ 添加文档", variant="primary")
                add_doc_msg = gr.Markdown("")
            
            gr.Markdown("---")
            gr.Markdown("### 🔍 搜索知识")
            search_input = gr.Textbox(label="搜索关键词", placeholder="如：颜体特点、行书笔法、永字八法", lines=1)
            search_btn = gr.Button("🔍 搜索", variant="primary")
            gr.Markdown("---")
            gr.Markdown("### 📑 文档列表")
            doc_list_dropdown = gr.Dropdown(label="选择文档", choices=[], interactive=True, type="value")
            refresh_doc_list_btn = gr.Button("🔄 刷新列表", size="sm")
            gr.Markdown("---")
            gr.Markdown("### 🗑️ 删除文档")
            delete_doc_btn = gr.Button("🗑️ 删除选中文档", variant="stop")
            delete_doc_msg = gr.Markdown("")
        
        with gr.Column(scale=2):
            gr.Markdown("### 📚 知识库")
            knowledge_result = gr.Markdown("")
    
    # 在所有组件创建后再绑定事件   
    add_doc_btn.click(
        fn=handle_add_document,
        inputs=[doc_file, manual_title, manual_content, video_select],
        outputs=[knowledge_result, add_doc_msg, doc_file, manual_title, manual_content, video_select, doc_list_dropdown]
    )
    
    search_btn.click(
        fn=handle_knowledge_search,
        inputs=[search_input],
        outputs=[knowledge_result]
    )
    
    delete_doc_btn.click(
        fn=handle_delete_knowledge_dropdown,
        inputs=[doc_list_dropdown],
        outputs=[knowledge_result, delete_doc_msg, doc_list_dropdown]
    )
    
    refresh_doc_list_btn.click(
        fn=lambda: gr.update(choices=get_doc_list_choices()),
        inputs=[],
        outputs=[doc_list_dropdown]
    )
    
    doc_list_dropdown.change(
        fn=show_selected_doc,
        inputs=[doc_list_dropdown],
        outputs=[knowledge_result]
    )
    
    return knowledge_result, video_select, doc_list_dropdown


def get_doc_list_choices():
    """获取文档下拉框选项列表"""
    docs = get_all_docs()
    choices = []
    for doc in docs:
        choice = f"{doc['id']}: {doc['title']}"
        choices.append(choice)
    return choices


def format_all_documents():
    """格式化显示所有文档"""
    docs = get_all_docs()
    if not docs:
        return "*暂无文档*"
    
    md = "### 📚 所有文档\n\n"
    for doc in docs:
        md += f"**{doc['id']}: {doc['title']}**\n"
        linked_videos = get_doc_videos(doc['id'])
        if linked_videos:
            md += "📺 "
            for v in linked_videos:
                md += f"{v['title']}, "
            md = md.rstrip(", ") + "\n"
        content_preview = doc['content'][:150]
        md += f"{content_preview}...\n\n"
    return md


def format_search_results(docs):
    """格式化显示搜索结果"""
    if not docs:
        return "*暂无匹配文档*"
    
    md = "### 🔍 搜索结果\n\n"
    for doc in docs:
        doc_id = doc.get('id', 'N/A')
        md += f"**{doc_id}: {doc['title']}**\n"
        content_preview = doc['content'][:150]
        md += f"{content_preview}...\n\n"
    return md


def handle_knowledge_search(keyword):
    """处理知识搜索"""
    if not keyword or not keyword.strip():
        # 空搜索显示所有文档
        docs = get_all_docs()
        print(f"[DEBUG] Empty search, showing all {len(docs)} documents")
        return format_all_documents()
    
    query = keyword.strip()
    print(f"[DEBUG] Searching for: '{query}'")
    
    # 先尝试获取所有文档数量
    try:
        all_docs = get_all_docs()
        print(f"[DEBUG] Total documents in DB: {len(all_docs)}")
        if all_docs:
            print(f"[DEBUG] Available titles: {[d['title'] for d in all_docs]}")
    except Exception as e:
        print(f"[DEBUG] Error getting all docs: {e}")
    
    # 执行搜索
    try:
        docs = search_knowledge(query, top_k=10)
        print(f"[DEBUG] search_knowledge returned {len(docs)} documents")
        for i, doc in enumerate(docs):
            print(f"[DEBUG] Result {i+1}: id={doc.get('id')}, title={doc.get('title')}, score={doc.get('score')}")
    except Exception as e:
        print(f"[DEBUG] search_knowledge error: {e}")
        # 失败时使用简单搜索
        try:
            from knowledge import simple_search
            docs = simple_search(query, top_k=10)
            print(f"[DEBUG] simple_search fallback returned {len(docs)} documents")
        except Exception as e2:
            print(f"[DEBUG] simple_search also failed: {e2}")
            return f"❌ 搜索失败: {str(e)}\n备援搜索也失败: {str(e2)}"
    
    if not docs:
        # 搜索无结果，显示所有文档让用户知道有什么
        docs = get_all_docs()
        if docs:
            md = "### 🔍 未找到匹配文档\n\n"
            md += "当前知识库中的文档有：\n\n"
            for doc in docs:
                md += f"- **{doc['title']}** (ID: {doc['id']})\n"
            md += "\n请尝试使用文档标题中的关键词搜索"
            return md
        else:
            return "*知识库暂无文档，请先上传文档*"
    
    return format_search_results(docs)
    
    if not docs:
        # 尝试简单搜索作为备援
        try:
            simple_results = simple_search(keyword.strip(), top_k=10)
            if simple_results:
                return format_search_results(simple_results)
        except:
            pass
        return "*未找到匹配文档，请尝试其他关键词*"
    
    return format_search_results(docs)


def handle_add_document(doc_file, title, content, video_id):
    from knowledge import get_db_path
    import sqlite3
    
    docs_dir = get_docs_dir()
    
    file_saved = False
    file_content = ""
    
    if doc_file:
        try:
            ext = os.path.splitext(doc_file.name)[1].lower()
            filename = f"{hash(doc_file.name)}_{os.path.basename(doc_file.name)}"
            dest_path = os.path.join(docs_dir, filename)
            
            if ext in ['.txt', '.md']:
                if hasattr(doc_file, 'read'):
                    content_bytes = doc_file.read()
                    with open(dest_path, 'wb') as f:
                        f.write(content_bytes)
                    file_content = content_bytes.decode('utf-8', errors='ignore')
                else:
                    with open(doc_file, 'r', encoding='utf-8') as f:
                        file_content = f.read()
                    import shutil
                    shutil.copy(doc_file, dest_path)
                
                file_saved = True
            elif ext == '.docx':
                try:
                    import docx
                    doc = docx.Document(doc_file)
                    paragraphs = [p.text for p in doc.paragraphs]
                    file_content = '\n'.join(paragraphs)
                    
                    if hasattr(doc_file, 'read'):
                        doc_file.seek(0)
                        content_bytes = doc_file.read()
                        with open(dest_path, 'wb') as f:
                            f.write(content_bytes)
                    else:
                        import shutil
                        shutil.copy(doc_file, dest_path)
                    
                    file_saved = True
                except ImportError:
                    return gr.update(), "⚠️ 请安装 python-docx 库: pip install python-docx", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
                except Exception as e:
                    return gr.update(), f"⚠️ DOCX文件读取失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
            elif ext == '.pdf':
                try:
                    import PyPDF2
                    if hasattr(doc_file, 'read'):
                        reader = PyPDF2.PdfReader(doc_file)
                    else:
                        reader = PyPDF2.PdfReader(doc_file)
                    for page in reader.pages:
                        file_content += page.extract_text() or ""
                    
                    if hasattr(doc_file, 'read'):
                        doc_file.seek(0)
                        content_bytes = doc_file.read()
                        with open(dest_path, 'wb') as f:
                            f.write(content_bytes)
                    else:
                        import shutil
                        shutil.copy(doc_file, dest_path)
                    
                    file_saved = True
                except ImportError:
                    return gr.update(), "⚠️ 请安装 PyPDF2 库: pip install PyPDF2", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
                except Exception as e:
                    return gr.update(), f"⚠️ PDF文件读取失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
            else:
                return gr.update(), "⚠️ 仅支持 txt/md/docx/pdf 文件", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
            
            if title and title.strip():
                title = title.strip()
            else:
                if doc_file:
                    title = os.path.basename(doc_file.name)
                else:
                    return gr.update(), "⚠️ 请填写文档标题", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
        except Exception as e:
            return gr.update(), f"⚠️ 文件读取失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
    
    if file_saved:
        final_content = file_content
    else:
        final_content = content
    
    if not title or not title.strip():
        return gr.update(), "⚠️ 请填写文档标题", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
    
    if not final_content or not final_content.strip():
        return gr.update(), "⚠️ 请填写文档内容或上传文件", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()
    
    try:
        doc_id = add_knowledge_doc(title.strip(), final_content.strip(), 'entry', None)
        
        if video_id:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute('''
                INSERT OR REPLACE INTO doc_videos (doc_id, video_id)
                VALUES (?, ?)
            ''', (doc_id, int(video_id)))
            conn.commit()
            conn.close()
        
        rebuild_index()
        
        # 更新文档列表和显示
        new_doc_list = get_doc_list_choices()
        updated_display = format_all_documents()
        
        return updated_display, "✅ 文档添加成功！", gr.update(), gr.update(value=""), gr.update(value=""), gr.update(value=None), gr.update(choices=new_doc_list, value=None)
    except Exception as e:
        return gr.update(), f"❌ 添加失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update(), gr.update()


def handle_delete_knowledge_dropdown(selected_index):
    """通过下拉框选择删除文档"""
    print(f"[DEBUG] handle_delete_knowledge_dropdown called with selected_index={selected_index}, type={type(selected_index)}")
    
    if selected_index is None or selected_index == '':
        print("[DEBUG] No selection")
        return "⚠️ 请先选择要删除的文档", "⚠️ 未选择", gr.update()
    
    try:
        docs = get_all_docs()
        print(f"[DEBUG] Total docs available: {len(docs)}")
        if not docs:
            return "暂无文档", "⚠️ 没有可删除的文档", gr.update()
        
        # selected_index 是 "uuid:标题" 格式，需要解析出 uuid
        doc_id = selected_index
        if ':' in str(selected_index):
            doc_id = str(selected_index).split(':')[0]
            print(f"[DEBUG] Extracted uuid from selection: {doc_id}")
        
        doc_info = None
        for doc in docs:
            if doc['id'] == doc_id:
                doc_info = doc
                break
        
        if not doc_info:
            print(f"[DEBUG] Document not found: {doc_id}")
            return "⚠️ 文档不存在", "⚠️ 未找到文档", gr.update()
        
        doc_title = doc_info['title']
        print(f"[DEBUG] Deleting doc_id={doc_id}, title='{doc_title}'")
        
        from knowledge import delete_knowledge_doc as db_delete_doc
        success = db_delete_doc(doc_id)
        print(f"[DEBUG] delete_knowledge_doc returned: {success}")
        
        if success:
            rebuild_index()  # 更新 lazyllm 向量索引
            print("[DEBUG] rebuild_index() called")
            # 刷新列表和显示
            new_doc_list = get_doc_list_choices()
            updated_display = format_all_documents()
            print(f"[DEBUG] Refreshed doc list: {len(new_doc_list)} items")
            return updated_display, "✅ 文档已删除", gr.update(choices=new_doc_list, value=None)
        else:
            print("[DEBUG] delete_knowledge_doc returned False")
            return gr.update(), "❌ 删除失败", gr.update()
    except Exception as e:
        print(f"[DEBUG] Exception in handle_delete_knowledge_dropdown: {e}")
        import traceback
        traceback.print_exc()
        return gr.update(), f"❌ 删除失败: {str(e)}", gr.update()


def get_video_choices():
    videos = get_all_videos(limit=50)
    videos = sort_videos_by_title(videos)
    choices = []
    for v in videos:
        v_id, v_title, _, _, _, _, _, v_views, _, _ = v
        choices.append((f"{v_title} (播放:{v_views})", str(v_id)))
    return choices


def handle_add_document(doc_file, title, content, video_id):
    file_content = ""
    
    if doc_file:
        try:
            ext = os.path.splitext(doc_file.name)[1].lower()
            
            if ext in ['.txt', '.md']:
                if hasattr(doc_file, 'read'):
                    content_bytes = doc_file.read()
                    file_content = content_bytes.decode('utf-8', errors='ignore')
                else:
                    with open(doc_file, 'r', encoding='utf-8') as f:
                        file_content = f.read()
            elif ext == '.docx':
                try:
                    import docx
                    doc = docx.Document(doc_file)
                    paragraphs = [p.text for p in doc.paragraphs]
                    file_content = '\n'.join(paragraphs)
                except ImportError:
                    return gr.update(), "⚠️ 请安装 python-docx 库: pip install python-docx", gr.update(), gr.update(), gr.update(), gr.update()
                except Exception as e:
                    return gr.update(), f"⚠️ DOCX文件读取失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update()
            elif ext == '.pdf':
                try:
                    import PyPDF2
                    if hasattr(doc_file, 'read'):
                        reader = PyPDF2.PdfReader(doc_file)
                    else:
                        reader = PyPDF2.PdfReader(doc_file)
                    for page in reader.pages:
                        file_content += page.extract_text() or ""
                except ImportError:
                    return gr.update(), "⚠️ 请安装 PyPDF2 库: pip install PyPDF2", gr.update(), gr.update(), gr.update(), gr.update()
                except Exception as e:
                    return gr.update(), f"⚠️ PDF文件读取失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update()
            else:
                return gr.update(), "⚠️ 仅支持 txt/md/docx/pdf 文件", gr.update(), gr.update(), gr.update(), gr.update()
            
            if title and title.strip():
                title = title.strip()
            else:
                title = os.path.basename(doc_file.name)
        except Exception as e:
            return gr.update(), f"⚠️ 文件读取失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update()
    
    if file_content:
        final_content = file_content
    else:
        final_content = content
    
    if not title or not title.strip():
        return gr.update(), "⚠️ 请填写文档标题", gr.update(), gr.update(), gr.update(), gr.update()
    
    if not final_content or not final_content.strip():
        return gr.update(), "⚠️ 请填写文档内容或上传文件", gr.update(), gr.update(), gr.update(), gr.update()
    
    try:
        doc_id = add_knowledge_doc(title.strip(), final_content.strip(), 'entry', None)
        
        if video_id:
            conn = sqlite3.connect(get_db_path())
            cursor = conn.cursor()
            cursor.execute('''
                INSERT OR REPLACE INTO doc_videos (doc_id, video_id)
                VALUES (?, ?)
            ''', (doc_id, int(video_id)))
            conn.commit()
            conn.close()
        
        rebuild_index()
        
        new_doc_list = get_doc_list_choices()
        updated_display = format_all_documents()
        
        return updated_display, "✅ 文档添加成功！", gr.update(), gr.update(value=""), gr.update(value=""), gr.update(value=None), gr.update(choices=new_doc_list, value=None)
    except Exception as e:
        return gr.update(), f"❌ 添加失败: {str(e)}", gr.update(), gr.update(), gr.update(), gr.update()


def handle_knowledge_search(keyword):
    if not keyword or not keyword.strip():
        docs = get_all_docs()
    else:
        try:
            docs = search_knowledge(keyword.strip(), top_k=10)
        except Exception as e:
            return f"❌ 搜索失败: {str(e)}"
    
    if not docs:
        return "*暂无相关知识*"
    
    md = "### 📚 搜索结果\n\n"
    for doc in docs:
        doc_id = doc.get('id', 'N/A')
        md += f"**{doc['title']}** (ID: {doc_id})\n"
        md += f"- {doc['content'][:200]}...\n\n"
    
    return md


def handle_delete_knowledge(doc_id):
    if not doc_id or not doc_id.strip():
        return "⚠️ 请输入要删除的文档ID", "⚠️ 请输入文档ID"
    
    try:
        from knowledge import delete_knowledge_doc as db_delete_doc, get_all_docs
        
        input_id = doc_id.strip()
        if ':' in input_id:
            input_id = input_id.split(':')[0]
        
        success = db_delete_doc(input_id)
        
        if success:
            docs = get_all_docs()
            md = "### 📚 知识库\n\n"
            for doc in docs:
                doc_id_val = doc['id']
                md += f"**{doc_id_val}: {doc['title']}**\n"
                md += f"- {doc['content'][:200]}...\n\n"
            
            return md, "✅ 文档已删除"
        else:
            return "❌ 删除失败", "❌ 删除失败"
    except Exception as e:
        return f"❌ 删除失败: {str(e)}", f"❌ 删除失败: {str(e)}"


def load_explain_chat_history():
    from knowledge import get_full_chat_history
    history = get_full_chat_history('default_user', limit=20)
    if not history:
        return [], None
    
    choices = [h['query'][:50] + ('...' if len(h['query']) > 50 else '') for h in history]
    default_val = choices[0] if choices else None
    return [], gr.update(value=default_val, choices=choices)


def update_history_dropdown():
    """仅更新下拉菜单，不更新 state"""
    from knowledge import get_full_chat_history
    history = get_full_chat_history('default_user', limit=20)
    if not history:
        return gr.update(value=None, choices=[])
    
    choices = [h['query'][:50] + ('...' if len(h['query']) > 50 else '') for h in history]
    default_val = choices[0] if choices else None
    return gr.update(value=default_val, choices=choices)


def handle_load_history(query, chat_history_state, video_dropdown):
    """加载选中的历史记录到右侧聊天窗口"""
    if not query:
        return gr.update(), gr.update(), chat_history_state, gr.update(value=None, choices=[]), gr.update(visible=False)
    
    from knowledge import get_full_chat_history
    history = get_full_chat_history('default_user', limit=20)
    history_item = next((h for h in history if h['query'] == query), None)
    
    if not history_item:
        return gr.update(), gr.update(), chat_history_state, gr.update(value=None, choices=[]), gr.update(visible=False)
    
    created_at = str(history_item['created_at'])
    answer = history_item.get('answer', '')
    
    chat_md = f"**🕐 {created_at}**\n\n**👤 用户：** {query}\n\n**🤖 回答：** {answer}"
    
    new_history = chat_history_state + [(query, answer)]
    
    return chat_md, gr.update(value=query), new_history, gr.update(value=None, choices=[]), gr.update(visible=False)

def build_explain_ui():
    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### 💡 智能问答")
            gr.Markdown("基于知识库进行RAG问答")
            
            with gr.Accordion("📎 上传参考资料(可选)", open=False):
                rag_file = gr.File(
                    label="选择文件 (txt, jpg, png, md, pdf)",
                    file_count="single",
                    file_types=[".txt", ".jpg", ".jpeg", ".png", ".md", ".pdf"]
                )
                file_content = gr.Textbox(
                    label="或直接输入内容",
                    placeholder="输入参考资料内容",
                    lines=4
                )
            
            gr.Markdown("### 🗂️ 对话历史（选择后继续聊天）")
            
            chat_history_dropdown = gr.Dropdown(
                label="选择历史问题",
                choices=[],
                interactive=True,
                allow_custom_value=True
            )
            
            with gr.Row():
                load_history_btn = gr.Button("📥 加载", variant="primary")
                delete_selected_btn = gr.Button("🗑️ 删除", variant="stop")
            
            new_chat_btn = gr.Button("➕ 新建对话", variant="primary")
            
            chat_history_state = gr.State(value=[])
            
            gr.Markdown("### 📺 相关视频")
            video_dropdown = gr.Dropdown(
                label="选择视频播放",
                choices=[],
                interactive=True
            )
            play_video_btn = gr.Button("▶️ 播放视频", variant="primary")
        
        with gr.Column(scale=2):
            gr.Markdown("### 💬 问答")
            
            chat_display = gr.Markdown("**💬 请在上方输入问题开始对话**", elem_id="chat-display")
            
            with gr.Row():
                query_input = gr.Textbox(
                    label="输入问题",
                    placeholder="请输入您的问题...",
                    lines=2
                )
            
            with gr.Row():
                submit_btn = gr.Button("📤 发送", variant="primary")
                stop_btn = gr.Button("⏹️ 停止", variant="stop")
            
            loading_indicator = gr.Markdown("")
    
    with gr.Row():
        video_player = gr.Video(label="视频播放", visible=False, height=800, width="100%")
    
    submit_btn.click(
        fn=handle_rag_chat,
        inputs=[query_input, rag_file, file_content, chat_history_state],
        outputs=[chat_display, query_input, chat_history_state, chat_history_dropdown, loading_indicator, video_dropdown]
    )
    
    new_chat_btn.click(
        fn=handle_new_chat,
        inputs=[],
        outputs=[chat_display, chat_history_state, chat_history_dropdown, video_dropdown, video_player]
    )
    
    load_history_btn.click(
        fn=handle_load_history,
        inputs=[chat_history_dropdown, chat_history_state, video_dropdown],
        outputs=[chat_display, query_input, chat_history_state, video_dropdown, video_player]
    )
    
    delete_selected_btn.click(
        fn=handle_delete_selected_history,
        inputs=[chat_history_dropdown, video_dropdown],
        outputs=[chat_history_dropdown, chat_display, video_dropdown, video_player]
    )
    
    stop_btn.click(
        fn=handle_stop_generation,
        inputs=[],
        outputs=[loading_indicator]
    )
    
    play_video_btn.click(
        fn=handle_play_related_video,
        inputs=[video_dropdown],
        outputs=[video_player]
    )
    
    return chat_history_dropdown, chat_history_state


def handle_rag_chat(query, rag_file, file_content, chat_history):
    import time
    import os
    from pathlib import Path
    
    if not query or not query.strip():
        return gr.update(), "", chat_history, gr.update(), "", gr.update(visible=False)
    
    has_file = False
    file_path = None
    
    if rag_file:
        try:
            if hasattr(rag_file, 'name'):
                file_path = rag_file.name
                has_file = True
            elif hasattr(rag_file, 'read'):
                has_file = True
        except:
            pass
    
    if file_path and has_file:
        ext = Path(file_path).suffix.lower()
        if ext in ['.jpg', '.jpeg', '.png']:
            from agent import get_agent
            vlm, _, _ = get_agent("VLM")
            analysis_prompt = f"请仔细分析这张图片中的内容，并从中提取与用户问题相关的关键信息。用户问题是：{query}\n\n请用简洁的语言总结图片中与问题相关的要点。"
            try:
                analysis_result = vlm(analysis_prompt, lazyllm_files=[file_path])
                analyzed_query = f"{query}\n\n图片分析结果: {analysis_result}"
                query = analyzed_query
            except Exception as e:
                print(f"VLM analysis error: {e}")
        elif ext in ['.txt', '.md']:
            try:
                if hasattr(rag_file, 'read'):
                    extra_content = rag_file.read().decode('utf-8', errors='ignore')
                else:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        extra_content = f.read()
                query = f"{query}\n\n参考资料: {extra_content}"
            except:
                pass
        elif ext == '.pdf':
            try:
                import PyPDF2
                pdf_text = ""
                if hasattr(rag_file, 'read'):
                    reader = PyPDF2.PdfReader(rag_file)
                else:
                    reader = PyPDF2.PdfReader(file_path)
                for page in reader.pages:
                    pdf_text += page.extract_text() or ""
                if pdf_text:
                    query = f"{query}\n\n参考资料: {pdf_text[:5000]}"
            except ImportError:
                query = f"{query}\n\n[PDF文件已上传，但PyPDF2未安装无法解析内容]"
            except Exception as e:
                print(f"PDF extract error: {e}")
    elif file_content and file_content.strip():
        query = f"{query}\n\n参考资料: {file_content}"
    
    timestamp = time.strftime("%Y-%m-%d %H:%M")
    chat_history = chat_history + [(timestamp, query.strip())]
    
    from knowledge import query_rag, reset_stop_flag, find_related_videos, search_knowledge, get_doc_videos
    reset_stop_flag()
    
    kb_results = search_knowledge(query, top_k=3)
    
    kb_content = ""
    linked_videos_from_kb = []
    if kb_results:
        kb_content = "**📚 知识库检索内容：**\n\n"
        for i, r in enumerate(kb_results, 1):
            kb_content += f"{i}. **{r.get('title', '文档')[:50]}**\n"
            kb_content += f"   {r.get('content', '')[:200]}...\n\n"
            
            doc_id = r.get('id')
            if doc_id:
                doc_videos = get_doc_videos(doc_id)
                linked_videos_from_kb.extend(doc_videos)
    else:
        kb_content = "*未从知识库中找到相关内容*\n\n"
    
    full_answer = query_rag(query)
    
    from knowledge import save_chat_history
    save_chat_history('default_user', query, full_answer)
    
    related_videos = find_related_videos(query, top_k=3)
    
    all_videos = linked_videos_from_kb + related_videos
    seen = set()
    unique_videos = []
    for v in all_videos:
        if v['id'] not in seen:
            seen.add(v['id'])
            unique_videos.append(v)
    
    video_links = ""
    if unique_videos:
        video_links = "\n\n---\n**📺 相关视频：**\n"
        for v in unique_videos:
            video_links += f"- [{v['title']}]({v['id']})\n"
    
    response = f"**🙋 你**: {query}\n\n{kb_content}---\n\n**🤖 AI**: {full_answer}{video_links}"
    
    if unique_videos:
        video_dropdown_choices = [(v['title'], str(v['id'])) for v in unique_videos]
        yield response, "", chat_history, update_history_dropdown(), "✅ 完成", gr.update(visible=True, choices=video_dropdown_choices, value=str(unique_videos[0]['id']))
    else:
        yield response, "", chat_history, update_history_dropdown(), "✅ 完成", gr.update(visible=False)


def handle_new_chat():
    from knowledge import reset_stop_flag
    reset_stop_flag()
    return "**💬 请在上方输入问题开始对话**", gr.update(), gr.update(value=None, choices=[]), gr.update(value=None, choices=[]), gr.update(visible=False)


def handle_delete_selected_history(selected_query, video_dropdown):
    """删除选中的历史记录"""
    from knowledge import delete_chat_history_record
    if not selected_query:
        return update_history_dropdown(), "⚠️ 请先选择要删除的记录", gr.update(value=None, choices=[]), gr.update(visible=False)
    
    try:
        delete_chat_history_record('default_user', selected_query)
        return update_history_dropdown(), f"✅ 已删除", gr.update(value=None, choices=[]), gr.update(visible=False)
    except Exception as e:
        return update_history_dropdown(), f"❌ 删除失败: {str(e)}", gr.update(value=None, choices=[]), gr.update(visible=False)


def handle_delete_history(chat_history, video_dropdown):
    from knowledge import reset_stop_flag
    reset_stop_flag()
    return [], update_history_dropdown(), "**历史已清空**", gr.update(value=None, choices=[]), gr.update(visible=False)


def handle_stop_generation():
    from knowledge import set_stop_flag
    set_stop_flag()
    return "⏹️ 已停止生成"


def handle_play_related_video(video_id):
    from video import get_video_by_id, update_video_views
    
    if not video_id:
        return gr.update(visible=False)
    
    try:
        video = get_video_by_id(int(video_id))
        if video:
            update_video_views(int(video_id))
            return gr.update(visible=True, value=video[3])
    except:
        pass
    
    return gr.update(visible=False)


def build_module_ui():
    build_video_ui()
    with gr.Row():
        gr.Markdown("---")
    build_knowledge_ui()
    with gr.Row():
        gr.Markdown("---")
    build_explain_ui()